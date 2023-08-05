

import threading
from tqdm import tqdm
import math
import re
import os
import sys
import time
import logging

try:
    import pdfplumber
    import chardet
    from docx import Document
except ImportError as e:
    logging.warning(str(e.args))


os.chdir(os.path.split(sys.path[0])[0])

from app.plugins.common import CounterLock, smart_database

logging.basicConfig()
logger = logging.getLogger()
logger.setLevel(logging.ERROR)


class Doc2Embbeding():

    def __init__(self, collection):

        self.docs = []
        self.collection = collection
        self.embedding_lock = CounterLock()
        
            
    def split_batch(self, data, batch_size=100):
        n = len(data)
        batch_num = int(math.ceil(n/batch_size))
        for idx in range(batch_num):
            batch = data[idx*batch_size:(idx+1)*batch_size]
            if not batch:
                continue
            yield list(zip(*batch))
    
    @staticmethod
    def data_preprocess(data):
        data = re.sub(r'！', "！\n", data)
        data = re.sub(r'：', "：\n", data)
        data = re.sub(r'。', "。\n", data)
        # data = re.sub(r'\r', "\n", data)
        data = re.sub(r'\n\n', "\n", data)
        data = re.sub(r"\n\s*\n", "\n", data)
        data = data.split("\n")
        data = list(filter(lambda x:True if len(x.strip())>10 else False, data))
        return data
    
    @staticmethod
    def read_file(file, file_path):
        if file.endswith(".pdf"):
            with pdfplumber.open(file_path) as pdf:
                data_list = []
                for page in pdf.pages:
                    data_list.append(page.extract_text())
                data = "\n".join(data_list)
        elif file.endswith(".txt"):
            with open(file_path, 'rb') as f:
                b = f.read()
                result = chardet.detect(b)
            with open(file_path, 'r', encoding=result['encoding']) as f:
                data = f.read()
        else:
            doc = Document(file_path)
            data = []
            for i in range(len(doc.paragraphs)):
                text = re.sub("\s", " ", doc.paragraphs[i].text) 
                if not text.strip():
                    continue
                data.append(text)
            for tb in doc.tables:
                for row in tb.rows:
                    content = ""
                    for cell in row.cells:
                        content += cell.text + "\t"
                    data.append(content)
            data = "\n".join(data)
        return data
    
    def make_index(self, source, texts, metadatas):
        data = {"collection":self.collection, "sr":source, "documents":texts, "metadatas":metadatas}
        thread = threading.Thread(target=smart_database.add, kwargs=data)
        thread.start()
        while self.embedding_lock.get_waiting_threads()>2:
            time.sleep(0.1)
    

    def get_docs(self, source_folder_path):
        user_name = os.path.split(source_folder_path)[-1]
        all_files=[]
        total_files = []
        for root, dirs, files in os.walk(source_folder_path):
            for file in files:
                all_files.append([root, file])
                
        for i in tqdm(range(len(all_files))):
            root, file=all_files[i]
            data = ""
            if file[-4:] not in [".txt", ".pdf"] and file[-5:] not in [".docx"]:
                continue
            try:
                file_path = os.path.join(root, file)
                total_files.append(file_path)
                data = self.read_file(file, file_path)
            except Exception as e:
                print("文件读取失败，当前文件已被跳过：",file,"。错误信息：",e)
            data = self.data_preprocess(data) 
            meta_data = [{"source": file}]*len(data)
            for elm in self.split_batch(list(zip(*[data, meta_data]))): 
                self.make_index(user_name, list(elm[0]), list(elm[1])) 
        return total_files

    def run(self, vectorstore_dir, source_folder=None):

        if not source_folder:
            source_folder = vectorstore_dir
        source_folder_path = os.path.join(os.getcwd(), source_folder)

        total_files = self.get_docs(source_folder_path)
        if len(total_files) == 0:
            logger.info(f"{source_folder_path} 目录没有数据")
            sys.exit(0)
        
        while self.embedding_lock.get_waiting_threads()>0:
            time.sleep(0.1)
            
        with self.embedding_lock:
            time.sleep(0.1)
            
        for file_name in total_files:
            os.remove(file_name)



if __name__ == "__main__":

    begin = time.time()


    vectorstore_dir="memory/txt"

    Doc2Embbeding = Doc2Embbeding(collection='doc2vec')
    Doc2Embbeding.run(vectorstore_dir=vectorstore_dir)

    print(f"used time : {time.time()-begin}")