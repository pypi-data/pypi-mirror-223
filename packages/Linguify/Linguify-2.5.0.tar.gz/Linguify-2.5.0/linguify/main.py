import email
import smtplib
import imaplib
import os
import openai
import docx
import datetime
from googleapiclient.discovery import build
from google.oauth2.credentials import Credentials
from google_auth_oauthlib.flow import InstalledAppFlow
from google.auth.transport.requests import Request
import base64
import mimetypes
import click
import sys
import json





''' Command-Line Interface (CLI) functions'''


@click.command()  # This is click syntax
def start():
    business_name = True
    while business_name:
        company = click.prompt('Enter Company Name')
        if len(company)>3:
            business_name = False
        else:
            click.echo("Company name must be over 3 characters")
    mail = False
    click.echo('Select the Email of your choice!')
    while not mail:
        email_choice = click.prompt('Gmail or Outlook')
        if email_choice.lower() == 'gmail':
            mail = True
            credentials_file = get_credentials_path()
            if not is_user_authenticated(): # Check if the user is already authenticated
                login(credentials_file)
            click.echo("You are logged in.")
            check_email_for_voicemails(credentials_file,company)
        elif email_choice.lower() == 'outlook':
            # Connect to the IMAP server and login
            mail = True
            completed = False
            while not completed:
                try:
                    email_address, email_password = get_email_info()
                    check_email_for_voicemails_OUTLOOK(email_address,email_password,company)
                    completed = True
                except Exception as e:
                    click.echo(e)
        else:
            click.echo('Invalid Email.')

    print(f"\nThank you for using Linguify!\n")


def get_email_info():
    try:
        priv_file_path = os.path.join(os.path.dirname(__file__), 'priv.json')
        with open(priv_file_path, 'r') as config_file:
            config_data = json.load(config_file)
            email_address = config_data.get("Email Address")
            email_password = config_data.get("Email Password")
    except:
        email_address = click.prompt("Enter Email Address")
        email_password = click.prompt("Enter Email Password (hidden)", hide_input=True)
    return email_address, email_password
    
def get_credentials_path():
    # Get the path of the current script (main.py) or module using __file__
    current_file_path = os.path.abspath(__file__)

    # Create the path to the "credentials" folder (located in the same directory as main.py)
    credentials_folder = os.path.join(os.path.dirname(current_file_path), "credentials")
    try:
        os.mkdir(credentials_folder)
    except:
        pass

    # Combine with the filename to get the full path to credentials.json
    credentials_path = os.path.join(credentials_folder, "credentials.json")
    with open(credentials_path,'w') as file:
        credits = {"installed":{"client_id":"233660003200-7s98f8683gt5r84p7k18fiurft53md9k.apps.googleusercontent.com","project_id":"linguify","auth_uri":"https://accounts.google.com/o/oauth2/auth","token_uri":"https://oauth2.googleapis.com/token","auth_provider_x509_cert_url":"https://www.googleapis.com/oauth2/v1/certs","client_secret":"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","redirect_uris":["http://localhost"]}}
        json.dump(credits, file)
    return credentials_path
    

def login(file_path):
    """Authorize the application to access Gmail."""
    
    # Get the Gmail API service to initiate the OAuth2 flow
    service = get_gmail_service(file_path)

    # Redirect the user to the Google sign-in page to grant permission
    flow = InstalledAppFlow.from_client_secrets_file(file_path, SCOPES)
    auth_url, _ = flow.authorization_url(prompt='consent', request_uri=None)


def is_user_authenticated():
    try:
        creds = Credentials.from_authorized_user_file('token.json', SCOPES)
        return creds.valid
    except:
        return False


""" Outlook preferences: """

def check_email_for_voicemails_OUTLOOK(email_address, email_password,company):
    global voicemailCount
    processed_content = {}
    # Connect to the IMAP server and login
    imap_server = "outlook.office365.com"
    connected = False
    while not connected:
        try:
            mail = imaplib.IMAP4_SSL(imap_server)
            connected = True
        except:
            imap_server = click.prompt('Enter IMAP Server')
    try:
        mail.login(email_address, email_password)
        mail.select("INBOX")
    except:
        raise Exception('Incorrect username or password!')


    priv_file_path = os.path.join(os.path.dirname(__file__), 'priv.json')
    with open(priv_file_path, 'w') as config_file:
        initial = {"Email Address": email_address, "Email Password": email_password}
        json.dump(initial, config_file)
    # Search for voicemail emails
    print("You are logged in!")
    print("Searching through emails for New Voicemails...")
    status, response = mail.search(None, "UNSEEN", 'SUBJECT "Voicemail"')
    voice_processed = False
    if status == "OK":
        email_ids = response[0].split()  
        if len(email_ids)>0:
            for email_id in email_ids:
                # Retrieve the email with the voicemail attachment
                status, email_data = mail.fetch(email_id, "(RFC822)")
                if status == "OK":
                    raw_email = email_data[0][1]
                    email_message = email.message_from_bytes(raw_email)
                    for part in email_message.walk():
                        voicemail_content,mime_type,format = process_email_attachment_OUTLOOK(part, True)
                        if voicemail_content != None:
                            # Mark the email as read
                            mail.store(email_id, "+FLAGS", "\\Seen")
                            content_hash = hash(voicemail_content)
                            if content_hash not in processed_content_hashes:
                                processed_content_hashes.add(content_hash)
                                processed_content[f'{content_hash}'] = [voicemail_content,mime_type,format]
        if processed_content:
            print(f"Found {len(processed_content)} New Voicemails!")
            for values in processed_content.values():
                try:
                    process_email_attachment_OUTLOOK(None,False,True,values[0],values[1],values[2])

                    voicemailCount-=1
                    if voicemailCount == 1:
                        print(f'Voicemail {voicemailCount}/{len(processed_content)} has been successfully processed!')
                    else:
                        print(f'Voicemails {voicemailCount}/{len(processed_content)} have been successfully processed!')
                    voicemailCount+=1
                    voice_processed = True
                except KeyError:
                    pass
                    
        if voice_processed:
            send_email_notification_OUTLOOK(company,voicemailCount-1)
            
        else:
            print('There are no New Voicemails!')

    # Close the connection to the IMAP server
    mail.logout()

def send_email_notification_OUTLOOK(company,transcriptions): 

    # Compose the email notification
    subject = f"Linguify Usage Notification: {company}"
    message = f"Company Name: {company}\nDate: {date} \nTranscriptions: {transcriptions} voicemail(s)\nTotal Cost: {transcriptions} x $4.59 = ${transcriptions*4.59}"  # Do you want to add GST

    # Connect to the SMTP server and login
    server = smtplib.SMTP("smtp.office365.com", 587)
    server.starttls()
    server.login("linguify@hotmail.com", 'L1ngu1fy@))$')

    # Send the email notification
    email_body = f"Subject: {subject}\n\n{message}"
    server.sendmail("linguify@hotmail.com", "abdul_khafagy2004@hotmail.com", email_body)  # (From, To, Message)

    # Close the connection to the SMTP server
    server.quit()




def process_email_attachment_OUTLOOK(part, verify=False,process=False,ready_attachment_data=None,mime=None,formats=None):
    # Extract the voicemail attachment from the email
    if not process:
        for format in VALID_FORMATS:
            mime_type = part.get_content_type()
            if format in mime_type:
                attachment_data = part.get_payload(decode=True)
                if verify:
                    return attachment_data,mime_type,format
        return None,None,None
    else:
        print("Converting Voicemail into audio file...")
        possible_extensions = mimetypes.guess_all_extensions(mime)
        
        if sys.platform.startswith('win'):
            desktop_path = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
        elif sys.platform.startswith('darwin') or sys.platform.startswith('linux'):
            desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        else:
            raise Exception("Unknown operating system.")
        new_folder_path = os.path.join(desktop_path, 'Linguify_Voicemails')
        try:
        # Create the new folder
            os.mkdir(new_folder_path)
            
        except FileExistsError: 
            pass
        if not os.path.isdir(f'{new_folder_path}/{date}/'):
            filename = os.path.join(new_folder_path,f'{date}')
            os.mkdir(filename)

        # Process the voicemail attachment
        print('AI Processing audio file...')
        if possible_extensions:
            for ext in possible_extensions:
                filename = f'{new_folder_path}/{date}/{voicemailCount}{ext}'
                with open(filename, "wb") as attachment_file:
                    attachment_file.write(ready_attachment_data)

                try:
                    process_audio_file(filename)
                    break
                except:
                    os.remove(filename)
        else:
            ext = f'.{formats}'
            filename = f'{new_folder_path}/{date}/{voicemailCount}{ext}'
            with open(filename, "wb") as attachment_file:
                    attachment_file.write(ready_attachment_data)
            try:
                process_audio_file(filename)
            except:
                raise Exception(f'formats = {formats}, mimetype = {mime}, extension = {ext}')
                



"""
These libraries provide essential functionalities for different aspects of Linguify: Gmail Edition. os allows you to interact with the file system,
imaplib facilitates communication with IMAP servers for email retrieval, smtplib enables sending email notifications through SMTP
servers, and email helps in working with email messages and their components.
"""
# Email account credentials and settings

SCOPES = ['https://www.googleapis.com/auth/gmail.readonly', 'https://www.googleapis.com/auth/gmail.modify']

date = datetime.datetime.now()
date = f'{date.strftime("%Y")}-{date.strftime("%m")}-{date.strftime("%d")}'

VALID_FORMATS = ['m4a', 'mp3', 'webm', 'mp4', 'mpga', 'wav', 'mpeg', 'ogg', 'oga', 'flac']
voicemailCount = 1



"""Gmail preferences:"""

def send_email_notification(credentials_file_path,company,transcriptions):
    # Compose the email notification
    
    message = email.message.EmailMessage()
    message['From'] = 'linguify@hotmail.com'
    message['To'] = 'abdul_khafagy2004@hotmail.com'
    message['Subject'] = f"Linguify Usage Notification: {company}"
    message.set_content(f"Company Name: {company}\nDate: {date} \nTranscriptions: {transcriptions} voicemail(s)\nTotal Cost: {transcriptions} x $4.59 = ${transcriptions*4.59}")  # Do you want to add GST


    '''# Get the Gmail API service'''
    service = get_gmail_service(credentials_file_path)
    

    # Create the email message
    raw_email = message.as_string().encode('utf-8')
    encoded_email = base64.urlsafe_b64encode(raw_email).decode('utf-8')
    email_message = {'raw': encoded_email}

    # Send the email notification using the Gmail API
    try:
        service.users().messages().send(userId='me', body=email_message).execute()
    
    except Exception as e:
        print(f"Error: {e}\nPlease contact the provider (abdul_khafagy2004@hotmail.com).\n")


def get_api_key():
    
    # Load the API key from the configuration file
    try:
        config_file_path = os.path.join(os.path.dirname(__file__), 'config.json')
        with open(config_file_path, 'r') as config_file:
            config_data = json.load(config_file)
            api_key = config_data.get("openai_api_key")
    except:
        config_file_path = os.path.join(os.path.dirname(__file__), 'config.json')
        with open(config_file_path, 'w') as config_file:
            initial = {"openai_api_key": "YOUR_OPENAI_API_KEY_HERE"}
            json.dump(initial, config_file)
        api_key = 'YOUR_OPENAI_API_KEY_HERE'


    if api_key == 'YOUR_OPENAI_API_KEY_HERE':
        api_key = input("Please enter your OpenAI API key: ")
        data = {"openai_api_key": api_key}
        with open(config_file_path, "w") as file:
            json.dump(data, file)

    # Initialize the OpenAI API with the retrieved key
    return api_key, config_file_path




def process_audio_file(file):
    global voicemailCount
    correct = False
    while not correct:
        try:
            openai.api_key,config_path = get_api_key()  
            openai.Model.list()
            correct = True
        except:
            os.remove(config_path)
            print('Incorrect OpenAI API key.')
    

    # Processing the voicemail
    audio_file = open(file, "rb")
    transcript = openai.Audio.transcribe(
        model="whisper-1", 
        file=audio_file, 
        prompt='Format it in point form.'
        )
    transcript=transcript['text']
    
    # Drafting the formatted response
    data=[ {"role": "system", "content": "Respond in this format and add what you may think are important and relevant comments for the clinic to know based on the voicemail in the additional comments section: Name of patient: \nAppointment Booking: \nRequested Procedure: \nDate and Time of appointment: \nPossible complications: \nAny further inquiries: \nAdditional Comments:"}]
    data.append({"role": "user", "content": transcript})
    completion = openai.ChatCompletion.create(
    model="gpt-3.5-turbo",
    messages=data
    )
    response = completion.choices[0].message['content']
    #data.append({"role": "assistant", "content": response})

    # Creating the Word document
    if sys.platform.startswith('win'):
        desktop_path = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
    elif sys.platform.startswith('darwin') or sys.platform.startswith('linux'):
        desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
    else:
        raise Exception("Unknown operating system.")
    new_folder_path = os.path.join(desktop_path, 'Linguify_Documents')
    print('Creating Word Document...')

    file_pathway = new_folder_path+f'/{date}.docx'
    try:
        os.mkdir(new_folder_path)  # Create the new folder

    except FileExistsError:
        pass  # Folder already exists
        
    try:
        document = docx.Document(file_pathway)
    except:
        document = docx.Document()
    
    header_text = f'Voicemail Number {voicemailCount}:\n\n'
    document.add_heading(header_text, level=1)
    voicemailCount +=1
    document.add_paragraph(f"{response}\n\n")
    document.add_heading(f'\t\t\t\t\tTranscription:', level=2)
    document.add_paragraph(f"\n{transcript}\n\n")
    document.save(file_pathway)  


'''------------------------------------------------------------------------------------------------------'''

def extract_audio_from_eml(part):
    # Get the 'parts' from the Gmail message
    current = part['parts']

    # Look for the audio part in the Gmail message parts
    loop = True
    
    while loop:
        for now in current:
            try:
                if now['parts']:
                    current = now['parts']
            except:
                loop = False
                MIMETYPE = now['mimeType']
                attachment_id= now['body']['attachmentId']
    for format in VALID_FORMATS:
        if format in MIMETYPE:
            break
    return attachment_id,MIMETYPE,format


def process_email_attachment(part, service, email_message,verify=False,process=False,ready_attachment_data=None,mime=None,formats=None):
    # Extract the voicemail attachment from the email
    found_format = False
    if not process:
        for format in VALID_FORMATS:
            mime_type = part['mimeType']
            
            if 'rfc822' in mime_type:
                attachment_id, mime_type, format= extract_audio_from_eml(part)
                found_format = True

            elif format in mime_type:
                attachment_id = part['body']['attachmentId']
                found_format = True

            if found_format:
                attachment = service.users().messages().attachments().get(
                userId='me',
                messageId=email_message['id'],
                id=attachment_id
                ).execute()

                # Decode the attachment data from base64
                attachment_data = base64.urlsafe_b64decode(attachment.get('data').encode('UTF-8'))
                if verify:
                    return attachment_data, mime_type,format
        return None,None,None
    else:
        print("Converting Voicemail into audio file...")
        # Use the MIME type to determine the file extension
        possible_extensions = mimetypes.guess_all_extensions(mime)
        
        if sys.platform.startswith('win'):
            desktop_path = os.path.join(os.path.join(os.environ['USERPROFILE']), 'Desktop')
        elif sys.platform.startswith('darwin') or sys.platform.startswith('linux'):
            desktop_path = os.path.join(os.path.expanduser("~"), "Desktop")
        else:
            raise Exception("Unknown operating system.")
        new_folder_path = os.path.join(desktop_path, 'Linguify_Voicemails')
        try:
        # Create the new folder
            os.mkdir(new_folder_path)
            
        except FileExistsError: 
            pass
        if not os.path.isdir(f'{new_folder_path}/{date}/'):
            filename = os.path.join(new_folder_path,f'{date}')
            os.mkdir(filename)

        # Process the voicemail attachment
        print('AI Processing audio file...')
        if possible_extensions:
            for ext in possible_extensions:
                filename = f'{new_folder_path}/{date}/{voicemailCount}{ext}'
                with open(filename, "wb") as attachment_file:
                    attachment_file.write(ready_attachment_data)

                try:
                    process_audio_file(filename)
                    break
                except:
                    os.remove(filename)
        else:
            ext = f'.{formats}'
            filename = f'{new_folder_path}/{date}/{voicemailCount}{ext}'
            with open(filename, "wb") as attachment_file:
                    attachment_file.write(ready_attachment_data)
            try:
                process_audio_file(filename)
            except:
                raise Exception(f'formats = {formats}, mimetype = {mime}, extension = {ext}')
        
            
    

'''------------------------------------------------------------------------------------------------------'''


def get_gmail_service(credentials_file_path):
    """Authenticate and create a Gmail API service."""
    creds = None
    # The file token.json stores the user's access and refresh tokens, and is
    # created automatically when the authorization flow completes for the first
    # time.
    if os.path.exists('token.json'):
        creds = Credentials.from_authorized_user_file('token.json', SCOPES)
    # If there are no (valid) credentials available, let the user log in.
    if not creds or not creds.valid:
        if creds and creds.expired and creds.refresh_token:
            creds.refresh(Request())
        else:
            flow = InstalledAppFlow.from_client_secrets_file(
                os.path.abspath(credentials_file_path), SCOPES)
            creds = flow.run_local_server(port=0)
        # Save the credentials for the next run
        with open('token.json', 'w') as token:
            token.write(creds.to_json())


        # Call the Gmail API
    service = build('gmail', 'v1', credentials=creds)
    return service


'''------------------------------------------------------------------------------------------------------'''
processed_content_hashes = set()

def check_email_for_voicemails(credentials_file_path,company):
    global voicemailCount
    processed_content = {}
    try:
        # Connect to the Gmail API service
        service = get_gmail_service(credentials_file_path)

        # Search for voicemail emails
        print("Searching through emails for New Voicemails...")
        results = service.users().messages().list(         # Number of emails available
            userId='me', q='is:unread subject:"Voicemail"'
        ).execute()

        messages = results.get('messages', []) # Number of voicemails available
        voice_processed = False
        if messages:
            for message in messages:
                msg = service.users().messages().get(userId='me', id=message['id']).execute()  # for part in msg['payload']['parts']
                for part in msg['payload']['parts']:
                    voicemail_content,mime_type,format = process_email_attachment(part, service,msg, True)
                    if voicemail_content != None:

                        # Mark the email as read
                        service.users().messages().modify(userId='me', id=message['id'], body={'removeLabelIds': ['UNREAD']}).execute()
                        content_hash = hash(voicemail_content)
                        if content_hash not in processed_content_hashes:
                            processed_content_hashes.add(content_hash)
                            processed_content[f'{content_hash}'] = [voicemail_content,mime_type,format]
        if processed_content:
            print(f"Found {len(processed_content)} New Voicemails!")  
            for values in processed_content.values():
                try:  # Checks if email has voicemail attachment
                    
                    process_email_attachment(None, None, None, False, True, values[0],values[1],values[2])

                    
                    voicemailCount-=1
                    if voicemailCount == 1:
                        print(f'Voicemail {voicemailCount}/{len(processed_content)} has been successfully processed!')
                    else:
                        print(f'Voicemails {voicemailCount}/{len(processed_content)} have been successfully processed!')
                    voicemailCount+=1
                    voice_processed = True
                    

                except KeyError:
                    pass
                    

        if voice_processed:
            send_email_notification(credentials_file_path,company,voicemailCount-1)
        else:
            print('There are no New Voicemails!')
            
    except Exception as e:
        print(f"Error occurred: {e}\nPlease contact the provider (abdul_khafagy2004@hotmail.com).\n")

'''------------------------------------------------------------------------------------------------------'''
